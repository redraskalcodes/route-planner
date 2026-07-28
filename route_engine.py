"""
Route engine — self-contained, no Excel dependency.
All inputs come from Google Calendar; outputs are Word docs.
"""

import pickle
import re
from datetime import date, datetime, timedelta
from pathlib import Path

import googlemaps
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build

SCOPES = ["https://www.googleapis.com/auth/calendar.readonly"]
LATE_PENALTY_PER_MIN = 1000  # heavily penalise missing a booked window


# ── Google Calendar ──────────────────────────────────────────────────────────

def get_calendar_service(credentials_path: Path, token_path: Path):
    creds = None
    if token_path.exists():
        with open(token_path, "rb") as f:
            creds = pickle.load(f)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(str(credentials_path), SCOPES)
            creds = flow.run_local_server(port=0)
        with open(token_path, "wb") as f:
            pickle.dump(creds, f)
    return build("calendar", "v3", credentials=creds)


def fetch_jobs(target_date: date, service, job_keyword: str, delivery_keyword: str) -> list[dict]:
    """
    Reads the day's calendar events and returns a list of job dicts.
    An event is included if its title contains job_keyword.
    It is classified as a delivery if the title also contains delivery_keyword,
    otherwise as a pickup.
    """
    day_start = datetime.combine(target_date, datetime.min.time()).isoformat() + "Z"
    day_end = datetime.combine(target_date + timedelta(days=1), datetime.min.time()).isoformat() + "Z"

    events = service.events().list(
        calendarId="primary",
        timeMin=day_start,
        timeMax=day_end,
        singleEvents=True,
        orderBy="startTime",
    ).execute().get("items", [])

    jobs = []
    for ev in events:
        title = ev.get("summary", "").lower()
        if job_keyword.lower() not in title:
            continue

        job_type = "delivery" if delivery_keyword.lower() in title else "pickup"

        m = re.search(r"x(\d+)", title)
        qty = int(m.group(1)) if m else 1

        address = (ev.get("location", "") or ev.get("description", "")).strip()
        if not address:
            continue

        start_str = ev["start"].get("dateTime", ev["start"].get("date"))
        end_str = ev["end"].get("dateTime", ev["end"].get("date"))
        slot_start = datetime.fromisoformat(start_str.replace("Z", "+00:00")).astimezone().replace(tzinfo=None)
        slot_end = datetime.fromisoformat(end_str.replace("Z", "+00:00")).astimezone().replace(tzinfo=None)

        jobs.append({
            "type": job_type,
            "address": address,
            "qty": qty,
            "slot_start": slot_start,
            "slot_end": slot_end,
        })

    return jobs


# ── Distance matrix ──────────────────────────────────────────────────────────

def build_matrix(gmaps, locations: list[str]) -> list[list[int]]:
    n = len(locations)
    matrix = [[0] * n for _ in range(n)]
    for i, origin in enumerate(locations):
        result = gmaps.distance_matrix([origin], locations, mode="driving", units="metric")
        for j, el in enumerate(result["rows"][0]["elements"]):
            matrix[i][j] = round(el["duration"]["value"] / 60) if el["status"] == "OK" else 9999
    return matrix


# ── Route scoring ────────────────────────────────────────────────────────────

def route_feasible_and_cost(order, matrix, jobs, start_time, service_time_min=15):
    current = 0
    current_time = start_time
    total_drive = 0
    penalty = 0
    feasible = True
    etas = []

    for idx in order:
        travel = matrix[current][idx + 1]
        arrival = current_time + timedelta(minutes=travel)
        job = jobs[idx]
        if arrival > job["slot_end"]:
            feasible = False
            penalty += int((arrival - job["slot_end"]).total_seconds() / 60) * LATE_PENALTY_PER_MIN
        visit_time = max(arrival, job["slot_start"])
        etas.append((visit_time, travel))
        total_drive += travel
        current_time = visit_time + timedelta(minutes=service_time_min)
        current = idx + 1

    total_drive += matrix[current][0]
    return feasible, total_drive + penalty, etas


# ── Route construction ───────────────────────────────────────────────────────

def nearest_neighbour_init(matrix, jobs, start_time, forced_first=None, service_time_min=15):
    n = len(jobs)
    unvisited = set(range(n))
    order = []
    current = 0
    current_time = start_time

    if forced_first is not None and forced_first in unvisited:
        idx = forced_first
        travel = matrix[current][idx + 1]
        arrival = current_time + timedelta(minutes=travel)
        order.append(idx)
        unvisited.remove(idx)
        current = idx + 1
        current_time = max(arrival, jobs[idx]["slot_start"]) + timedelta(minutes=service_time_min)

    while unvisited:
        candidates = [
            (matrix[current][idx + 1], idx, current_time + timedelta(minutes=matrix[current][idx + 1]))
            for idx in unvisited
            if current_time + timedelta(minutes=matrix[current][idx + 1]) <= jobs[idx]["slot_end"]
        ]
        if candidates:
            travel, idx, arrival = min(candidates, key=lambda t: t[0])
        else:
            idx = min(unvisited, key=lambda i: jobs[i]["slot_end"])
            travel = matrix[current][idx + 1]
            arrival = current_time + timedelta(minutes=travel)

        order.append(idx)
        unvisited.remove(idx)
        current = idx + 1
        current_time = max(arrival, jobs[idx]["slot_start"]) + timedelta(minutes=service_time_min)

    return order


def local_search_improve(order, matrix, jobs, start_time, max_iters=300, service_time_min=15):
    _, best_score, _ = route_feasible_and_cost(order, matrix, jobs, start_time, service_time_min)
    improved = True
    iters = 0
    n = len(order)

    while improved and iters < max_iters:
        improved = False
        iters += 1

        for i in range(n - 1):
            for j in range(i + 1, n):
                cand = order[:i] + order[i:j + 1][::-1] + order[j + 1:]
                _, score, _ = route_feasible_and_cost(cand, matrix, jobs, start_time, service_time_min)
                if score < best_score:
                    order, best_score = cand, score
                    improved = True

        for i in range(n):
            stop = order[i]
            rest = order[:i] + order[i + 1:]
            for pos in range(len(rest) + 1):
                cand = rest[:pos] + [stop] + rest[pos:]
                if cand == order:
                    continue
                _, score, _ = route_feasible_and_cost(cand, matrix, jobs, start_time, service_time_min)
                if score < best_score:
                    order, best_score = cand, score
                    improved = True

    return order, best_score


def multi_start_optimise(matrix, jobs, start_time, service_time_min=15):
    best_order, best_score = None, float("inf")
    for first in range(len(jobs)):
        init = nearest_neighbour_init(matrix, jobs, start_time, forced_first=first, service_time_min=service_time_min)
        order, score = local_search_improve(init, matrix, jobs, start_time, service_time_min=service_time_min)
        if score < best_score:
            best_order, best_score = order, score
    return best_order, best_score


def optimise_route(gmaps, jobs: list[dict], start_address: str, service_time_min: int = 15) -> list[dict]:
    locations = [start_address] + [j["address"] for j in jobs]
    matrix = build_matrix(gmaps, locations)

    start_time = min(j["slot_start"] for j in jobs)
    best_order, _ = multi_start_optimise(matrix, jobs, start_time, service_time_min)

    first_idx = best_order[0]
    travel_to_first = matrix[0][first_idx + 1]
    departure_time = jobs[first_idx]["slot_start"] - timedelta(minutes=travel_to_first)

    feasible, drive_total, etas = route_feasible_and_cost(best_order, matrix, jobs, departure_time, service_time_min)

    ordered = []
    for pos, idx in enumerate(best_order):
        job = jobs[idx]
        visit_time, travel = etas[pos]
        job["travel_min"] = travel
        job["eta"] = visit_time
        ordered.append(job)

    last_idx = best_order[-1]
    return_travel = matrix[last_idx + 1][0]
    return_eta = etas[-1][0] + timedelta(minutes=service_time_min + return_travel)
    ordered.append({
        "type": "office",
        "address": start_address,
        "qty": 0,
        "travel_min": return_travel,
        "eta": return_eta,
        "slot_start": return_eta,
        "slot_end": return_eta,
    })

    print(f"  → {drive_total} min total driving (feasible: {feasible})")
    return ordered


# ── Helpers ──────────────────────────────────────────────────────────────────

def format_time(dt: datetime) -> str:
    h, m = dt.hour, dt.minute
    period = "am" if h < 12 else "pm"
    h12 = h if h <= 12 else h - 12
    if h12 == 0:
        h12 = 12
    return f"{h12}{'%02d' % m if m else ''}{period}" if m else f"{h12}{period}"


def generate_route_doc(target_date: date, ordered_jobs: list[dict], output_dir: Path) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    day_str = target_date.strftime("%A, ") + str(int(target_date.strftime("%d"))) + target_date.strftime(" %B %Y")
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(f"Route Plan — {day_str}")
    run.bold = True
    run.font.size = Pt(13)
    doc.add_paragraph()

    for job in ordered_jobs:
        travel = job.get("travel_min", "?")
        eta = format_time(job["eta"]) if "eta" in job else "?"

        if job["type"] == "office":
            p = doc.add_paragraph(f"{job['address']}  ({travel} min)  {eta}")
            p.runs[0].bold = True
            continue

        slot_label = f"[booked {format_time(job['slot_start'])}–{format_time(job['slot_end'])}]"
        qty = job["qty"]

        line = f"{job['address']}  ({travel} min)  {eta}  {slot_label}"
        if job["type"] == "delivery":
            line += f"    ⇒ Del {qty} bag{'s' if qty > 1 else ''}"

        doc.add_paragraph(line)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"Generated {day_str}").italic = True
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / f"{target_date.strftime('%d-%m-%y')} Route Plan.docx"
    doc.save(str(out_path))
    return out_path
